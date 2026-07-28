from pathlib import Path


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return _extract_docx(path)
    elif ext == ".pdf":
        return _extract_pdf(path)
    elif ext in (".md", ".txt"):
        return path.read_text(encoding="utf-8", errors="replace")
    return ""


def _extract_docx(path: Path) -> str:
    import docx  # python-docx

    doc = docx.Document(str(path))
    parts: list[str] = []
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                parts.append(para.text)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    for section in doc.sections:
        for para in section.footer.paragraphs:
            if para.text.strip():
                parts.append(para.text)
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            # Lower x_tolerance (pdfplumber default is 3) to avoid gluing
            # adjacent words together when a PDF's character spacing is tight.
            text = page.extract_text(x_tolerance=1)
            if text:
                parts.append(text)
    return "\n".join(parts)
