from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

from resume_mcp_server.extractor import extract_text

_MD_CONTENT = """\
# Jane Doe

jane@example.com | 503-555-1234

## Summary

Experienced software engineer.

## Skills

Python, Go, Rust
"""

_TXT_CONTENT = "Jane Doe\njane@example.com\n503-555-1234\n\nSkills\nPython, Go"


class TestExtractMarkdown(unittest.TestCase):
    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self.path = Path(self._tmpdir.name) / "resume.md"
        self.path.write_text(_MD_CONTENT, encoding="utf-8")

    def tearDown(self):
        self._tmpdir.cleanup()

    def test_returns_string(self):
        result = extract_text(self.path)
        self.assertIsInstance(result, str)

    def test_not_empty(self):
        result = extract_text(self.path)
        self.assertGreater(len(result), 0)

    def test_preserves_content(self):
        result = extract_text(self.path)
        self.assertIn("Jane Doe", result)
        self.assertIn("Python", result)

    def test_preserves_markdown_headers(self):
        result = extract_text(self.path)
        self.assertIn("##", result)


class TestExtractTxt(unittest.TestCase):
    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self.path = Path(self._tmpdir.name) / "resume.txt"
        self.path.write_text(_TXT_CONTENT, encoding="utf-8")

    def tearDown(self):
        self._tmpdir.cleanup()

    def test_returns_string(self):
        result = extract_text(self.path)
        self.assertIsInstance(result, str)

    def test_preserves_content(self):
        result = extract_text(self.path)
        self.assertIn("Jane Doe", result)
        self.assertIn("Python", result)


class TestExtractUnsupportedExtension(unittest.TestCase):
    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self.path = Path(self._tmpdir.name) / "data.csv"
        self.path.write_text("col1,col2\nval1,val2")

    def tearDown(self):
        self._tmpdir.cleanup()

    def test_unsupported_returns_empty_string(self):
        result = extract_text(self.path)
        self.assertEqual(result, "")


class TestExtractUnicodeContent(unittest.TestCase):
    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self.path = Path(self._tmpdir.name) / "resume.txt"
        self.path.write_text("Résumé: María García\nSkills: Python", encoding="utf-8")

    def tearDown(self):
        self._tmpdir.cleanup()

    def test_unicode_roundtrip(self):
        result = extract_text(self.path)
        self.assertIn("María", result)
        self.assertIn("García", result)


class TestExtractEmptyFile(unittest.TestCase):
    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self.md_path = Path(self._tmpdir.name) / "empty.md"
        self.md_path.write_text("")
        self.txt_path = Path(self._tmpdir.name) / "empty.txt"
        self.txt_path.write_text("")

    def tearDown(self):
        self._tmpdir.cleanup()

    def test_empty_md_returns_empty_string(self):
        self.assertEqual(extract_text(self.md_path), "")

    def test_empty_txt_returns_empty_string(self):
        self.assertEqual(extract_text(self.txt_path), "")


class TestExtractDocx(unittest.TestCase):
    def setUp(self):
        import docx

        self._tmpdir = tempfile.TemporaryDirectory()
        self.path = Path(self._tmpdir.name) / "resume.docx"

        doc = docx.Document()
        doc.add_paragraph("Jane Doe")
        doc.add_paragraph("jane@example.com")
        doc.add_paragraph("   ")  # whitespace-only paragraph, should be skipped

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Python"
        table.cell(0, 1).text = "Expert"
        table.cell(1, 0).text = ""
        table.cell(1, 1).text = "  "  # whitespace-only cell, should be filtered

        doc.save(str(self.path))

    def tearDown(self):
        self._tmpdir.cleanup()

    def test_returns_string(self):
        result = extract_text(self.path)
        self.assertIsInstance(result, str)

    def test_preserves_paragraph_content(self):
        result = extract_text(self.path)
        self.assertIn("Jane Doe", result)
        self.assertIn("jane@example.com", result)

    def test_blank_paragraph_not_included(self):
        result = extract_text(self.path)
        lines = result.split("\n")
        self.assertNotIn("", lines)

    def test_table_row_cells_joined(self):
        result = extract_text(self.path)
        self.assertIn("Python  Expert", result)

    def test_table_row_with_all_blank_cells_omitted(self):
        result = extract_text(self.path)
        # Second table row has one empty and one whitespace-only cell,
        # so it should contribute no line to the output at all.
        lines = result.split("\n")
        self.assertEqual(len(lines), 3)  # "Jane Doe", "jane@example.com", "Python  Expert"


class TestExtractPdf(unittest.TestCase):
    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self.path = Path(self._tmpdir.name) / "resume.pdf"
        self.path.write_bytes(b"%PDF-1.4 fake content")

    def tearDown(self):
        self._tmpdir.cleanup()

    def _mock_pdf(self, page_texts):
        pages = []
        for text in page_texts:
            page = MagicMock()
            page.extract_text.return_value = text
            pages.append(page)
        mock_pdf = MagicMock()
        mock_pdf.pages = pages
        mock_pdf.__enter__.return_value = mock_pdf
        mock_pdf.__exit__.return_value = False
        return mock_pdf

    def test_returns_string(self):
        with patch("pdfplumber.open", return_value=self._mock_pdf(["Jane Doe"])):
            result = extract_text(self.path)
        self.assertIsInstance(result, str)

    def test_joins_page_text_with_newline(self):
        with patch("pdfplumber.open", return_value=self._mock_pdf(["Page one", "Page two"])):
            result = extract_text(self.path)
        self.assertEqual(result, "Page one\nPage two")

    def test_pages_with_no_extractable_text_are_skipped(self):
        with patch("pdfplumber.open", return_value=self._mock_pdf(["Page one", None, "Page three"])):
            result = extract_text(self.path)
        self.assertEqual(result, "Page one\nPage three")

    def test_all_pages_empty_returns_empty_string(self):
        with patch("pdfplumber.open", return_value=self._mock_pdf([None, None])):
            result = extract_text(self.path)
        self.assertEqual(result, "")


if __name__ == "__main__":
    unittest.main()
